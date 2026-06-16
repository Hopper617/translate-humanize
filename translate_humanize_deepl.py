from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import urllib.request
import urllib.parse
import json
import os
import sys
import argparse

API_KEY = os.environ.get("DEEPL_API_KEY")
ENDPOINT = os.environ.get("DEEPL_ENDPOINT", "https://api-free.deepl.com/v2/translate")


def deepl_translate(text: str, source_lang: str, target_lang: str) -> str:
    import time
    import urllib.error

    data = urllib.parse.urlencode({
        "text": text,
        "source_lang": source_lang,
        "target_lang": target_lang,
    }).encode()
    req = urllib.request.Request(
        ENDPOINT,
        data=data,
        headers={"Authorization": f"DeepL-Auth-Key {API_KEY}"},
        method="POST",
    )
    last_err = None
    for attempt in range(3):
        try:
            with urllib.request.urlopen(req, timeout=60) as resp:
                result = json.loads(resp.read().decode("utf-8"))
                return result["translations"][0]["text"]
        except (urllib.error.URLError, TimeoutError) as e:
            last_err = e
            wait = 2 ** attempt
            print(f"Translation {source_lang}->{target_lang} failed (attempt {attempt + 1}/3): {e}. Retrying in {wait}s...")
            time.sleep(wait)
    raise last_err


def read_sections_from_docx(path: str) -> dict:
    doc = Document(path)
    sections = {}
    current = None
    buffer = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in [
            "Original",
            "Translated English",
            "Japanese",
            "Korean",
            "Spanish",
            "Portuguese",
        ]:
            if current:
                sections[current] = "\n".join(buffer).strip()
            current = text
            buffer = []
        else:
            buffer.append(text)
    if current:
        sections[current] = "\n".join(buffer).strip()
    return sections


def main():
    parser = argparse.ArgumentParser(
        description="Generate a DeepL round-trip translation chain for agent-driven humanization."
    )
    parser.add_argument("--input", "-i", default="input.docx", help="Input docx containing an 'Original' section.")
    parser.add_argument("--output", "-o", default="translate_humanize_output.docx", help="Output docx path.")
    parser.add_argument("--deepl-key", default=None, help="DeepL API key (or set DEEPL_API_KEY env var).")
    parser.add_argument("--deepl-endpoint", default=None, help="DeepL API endpoint.")
    args = parser.parse_args()

    global API_KEY, ENDPOINT
    if args.deepl_key:
        API_KEY = args.deepl_key
    if args.deepl_endpoint:
        ENDPOINT = args.deepl_endpoint

    if not API_KEY:
        print("Error: DEEPL_API_KEY not set. Provide it via --deepl-key or the DEEPL_API_KEY environment variable.")
        sys.exit(1)

    sections = read_sections_from_docx(args.input)
    original = sections.get("Original", "").strip()
    if not original:
        print(f"Error: 'Original' section not found in {args.input}")
        sys.exit(1)

    print("Translating Original -> Japanese...")
    ja = deepl_translate(original, "EN", "JA")
    print("Translating Japanese -> Korean...")
    ko = deepl_translate(ja, "JA", "KO")
    print("Translating Korean -> Spanish...")
    es = deepl_translate(ko, "KO", "ES")
    print("Translating Spanish -> Portuguese...")
    pt = deepl_translate(es, "ES", "PT")
    print("Translating Portuguese -> English (Translated English)...")
    translated_en = deepl_translate(pt, "PT", "EN")

    texts = {
        "Original": original,
        "Japanese": ja,
        "Korean": ko,
        "Spanish": es,
        "Portuguese": pt,
        "Translated English": translated_en,
    }

    doc = Document()
    title = doc.add_heading("translate-humanize Translation Chain Output", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for label, text in texts.items():
        doc.add_heading(label, level=1)
        for para in text.split("\n"):
            p = doc.add_paragraph(para)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.size = Pt(11)
        doc.add_paragraph()

    doc.save(args.output)
    print(f"Saved translation-chain docx to: {args.output}")
    print(f"Total versions included: {len(texts)}")


if __name__ == "__main__":
    main()
