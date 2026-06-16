from docx import Document

paragraph = (
    "This controlled comparison supports the stronger claim that language is not just a useful extra feature, "
    "but rather a structural prerequisite for flexible task composition in robotic manipulation."
)

doc = Document()
doc.add_heading("Original", level=1)
doc.add_paragraph(paragraph)
doc.save("examples/sample_input.docx")
print("Created examples/sample_input.docx")
